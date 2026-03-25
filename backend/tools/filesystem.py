"""File system tools: read, write, edit."""

from pathlib import Path
from typing import Any

from backend.tools.base import Tool


class ReadFileTool(Tool):
    """Tool to read file contents. Supports txt, pdf, docx, xlsx, xls files."""
    
    @property
    def name(self) -> str:
        return "read_file"
    
    @property
    def description(self) -> str:
        return "Read the contents of a file at the given path. Supports txt, pdf, docx, xlsx, xls formats."
    
    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "The file path to read"
                }
            },
            "required": ["path"]
        }
    
    def _read_pdf(self, file_path: Path) -> str:
        try:
            import pypdf
            reader = pypdf.PdfReader(str(file_path))
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            return "\n\n".join(text_parts)
        except ImportError:
            return "Error: pypdf is not installed. Install it with: pip install pypdf"
    
    def _read_docx(self, file_path: Path) -> str:
        try:
            from docx import Document
            doc = Document(str(file_path))
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    table_text.append(" | ".join(row_text))
                if table_text:
                    text_parts.append("\n".join(table_text))
            return "\n\n".join(text_parts)
        except ImportError:
            return "Error: python-docx is not installed. Install it with: pip install python-docx"
    
    def _read_excel(self, file_path: Path) -> str:
        try:
            import openpyxl
            wb = openpyxl.load_workbook(str(file_path), read_only=True, data_only=True)
            all_sheets = []
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                rows = []
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        row_str = " | ".join(str(cell) if cell is not None else "" for cell in row)
                        rows.append(row_str)
                if rows:
                    all_sheets.append(f"=== Sheet: {sheet_name} ===\n" + "\n".join(rows))
            wb.close()
            return "\n\n".join(all_sheets)
        except ImportError:
            return "Error: openpyxl is not installed. Install it with: pip install openpyxl"
    
    async def execute(self, path: str, **kwargs: Any) -> str:
        from backend.utils.helpers import get_workspace_path
        file_path = Path(path).expanduser()
        if not file_path.is_absolute():
            file_path = get_workspace_path() / file_path
        try:
            if not file_path.exists():
                return f"Error: File not found: {path}"
            if not file_path.is_file():
                return f"Error: Not a file: {path}"
            
            suffix = file_path.suffix.lower()
            
            if suffix == ".pdf":
                return self._read_pdf(file_path)
            elif suffix == ".docx":
                return self._read_docx(file_path)
            elif suffix in (".xlsx", ".xls"):
                return self._read_excel(file_path)
            else:
                content = file_path.read_text(encoding="utf-8")
                return content
        except PermissionError:
            return f"Error: Permission denied: {path}"
        except Exception as e:
            return f"Error reading file: {str(e)}"


class WriteFileTool(Tool):
    """Tool to write content to a file."""
    
    @property
    def name(self) -> str:
        return "write_file"
    
    @property
    def description(self) -> str:
        return "Write content to a file at the given path. Creates parent directories if needed."
    
    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "The file path to write to"
                },
                "content": {
                    "type": "string",
                    "description": "The content to write"
                }
            },
            "required": ["path", "content"]
        }
    
    async def execute(self, path: str, content: str, **kwargs: Any) -> str:
        from backend.utils.helpers import get_workspace_path
        file_path = Path(path).expanduser()
        if not file_path.is_absolute():
            file_path = get_workspace_path() / file_path
        try:
            file_path.parent.mkdir(parents=True, exist_ok=True)
            file_path.write_text(content, encoding="utf-8")
            return f"Successfully wrote {len(content)} bytes to {path}"
        except PermissionError:
            return f"Error: Permission denied: {path}"
        except Exception as e:
            return f"Error writing file: {str(e)}"


class EditFileTool(Tool):
    """Tool to edit a file by replacing text."""
    
    @property
    def name(self) -> str:
        return "edit_file"
    
    @property
    def description(self) -> str:
        return "Edit a file by replacing old_text with new_text. The old_text must exist exactly in the file."
    
    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "The file path to edit"
                },
                "old_text": {
                    "type": "string",
                    "description": "The exact text to find and replace"
                },
                "new_text": {
                    "type": "string",
                    "description": "The text to replace with"
                }
            },
            "required": ["path", "old_text", "new_text"]
        }
    
    async def execute(self, path: str, old_text: str, new_text: str, **kwargs: Any) -> str:
        from backend.utils.helpers import get_workspace_path
        file_path = Path(path).expanduser()
        if not file_path.is_absolute():
            file_path = get_workspace_path() / file_path
        try:
            if not file_path.exists():
                return f"Error: File not found: {path}"
            
            content = file_path.read_text(encoding="utf-8")
            
            if old_text not in content:
                return f"Error: old_text not found in file. Make sure it matches exactly."
            
            # Count occurrences
            count = content.count(old_text)
            if count > 1:
                return f"Warning: old_text appears {count} times. Please provide more context to make it unique."
            
            new_content = content.replace(old_text, new_text, 1)
            file_path.write_text(new_content, encoding="utf-8")
            
            return f"Successfully edited {path}"
        except PermissionError:
            return f"Error: Permission denied: {path}"
        except Exception as e:
            return f"Error editing file: {str(e)}"


class ListDirTool(Tool):
    """Tool to list directory contents."""
    
    @property
    def name(self) -> str:
        return "list_dir"
    
    @property
    def description(self) -> str:
        return "List the contents of a directory."
    
    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "The directory path to list (defaults to workspace root)"
                }
            },
            "required": []
        }
    
    async def execute(self, path: str = ".", **kwargs: Any) -> str:
        from backend.utils.helpers import get_workspace_path
        if not path:
            dir_path = get_workspace_path()
        else:
            dir_path = Path(path).expanduser()
            if not dir_path.is_absolute():
                dir_path = get_workspace_path() / dir_path
        try:
            if not dir_path.exists():
                return f"Error: Directory not found: {path}"
            if not dir_path.is_dir():
                return f"Error: Not a directory: {path}"
            
            items = []
            for item in sorted(dir_path.iterdir()):
                prefix = "📁 " if item.is_dir() else "📄 "
                items.append(f"{prefix}{item.name}")
            
            if not items:
                return f"Directory {path} is empty"
            
            return "\n".join(items)
        except PermissionError:
            return f"Error: Permission denied: {path}"
        except Exception as e:
            return f"Error listing directory: {str(e)}"
